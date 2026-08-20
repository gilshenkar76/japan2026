import csv
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

title = doc.add_heading('Japan 2026 – Full Trip Summary', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Trip: November 12–29, 2026 | Group Trip').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

headers = ['Date', 'City/Area', 'Time', 'Attraction / Event', 'Type', 'Hotel']
col_widths = [Cm(2.2), Cm(3.5), Cm(2.5), Cm(5.5), Cm(3.2), Cm(4.5)]

with open('trip_summary.csv', encoding='utf-8') as f:
    rows = list(csv.DictReader(f))

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Header row
hdr = table.rows[0]
for cell, h in zip(hdr.cells, headers):
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '2E4057')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

TYPE_COLORS = {
    'Flight': 'D6EAF8', 'Transit': 'EBF5FB', 'Arrival': 'D5F5E3',
    'Temple/Shrine': 'FADBD8', 'Pagoda/Viewpoint': 'FADBD8', 'Pagoda': 'FADBD8',
    'Temple (Zen)': 'FADBD8', 'Temple/Garden': 'FADBD8', 'Temple/Market': 'FADBD8',
    'Historic Castle': 'F9EBEA', 'Historic District': 'FDEBD0', 'Historic Street': 'FDEBD0',
    'Traditional Village': 'FEF9E7',
    'Museum': 'E8DAEF', 'Science Museum': 'E8DAEF',
    'Digital Art': 'D7BDE2', 'Aquarium': 'D6EAF8',
    'Park/Garden': 'D5F5E3', 'Nature/Lake': 'D5F5E3', 'Nature/Park': 'D5F5E3',
    'Nature': 'D5F5E3', 'Urban Park': 'D5F5E3', 'Rooftop/Garden': 'D5F5E3',
    'Volcanic Valley': 'FDEBD0', 'Ropeway': 'EBF5FB', 'Cruise': 'D6EAF8',
    'Viewpoint': 'FEF9E7',
    'Food': 'FEF9E7', 'Food Market': 'FEF9E7', 'Food Market/District': 'FEF9E7',
    'Food/Alley': 'FEF9E7', 'Food/Waterfront': 'FEF9E7', 'Cafe': 'FEF9E7',
    'Shopping': 'FCF3CF', 'Shopping/Evening': 'FCF3CF', 'Shopping/Rooftop': 'FCF3CF',
    'Electronics/Anime': 'FCF3CF', 'Showroom': 'FCF3CF',
    'Bar': 'F9EBEA', 'Alley/Bars': 'F9EBEA', 'Nightlife': 'F9EBEA',
    'Nightlife District': 'F9EBEA', 'Nightlife/Cruise': 'F9EBEA',
    'Hotel': 'F2F3F4', 'Hotel/Onsen': 'F2F3F4', 'Logistics': 'FDFEFE',
    'Train': 'EBF5FB', 'Meeting': 'FDFEFE',
}

for row in rows:
    tr = table.add_row()
    vals = [row['Date'], row['City/Area'], row['Time'], row['Attraction / Event'], row['Type'], row['Hotel']]
    color = TYPE_COLORS.get(row['Type'], 'FFFFFF')
    for cell, val in zip(tr.cells, vals):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        if color != 'FFFFFF':
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color)
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

doc.save('trip_summary.docx')
print('Done: trip_summary.docx created')
